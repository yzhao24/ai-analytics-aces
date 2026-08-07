"""
Group Assignment 2 — Evaluation & Measurement, as a Word document.

    python3 build_assignment2_docx.py

Keeps the team's numbered specification structure for Part A and adds Parts
B-D. Every figure comes from score_rubric.py, stability_test.py and
usability_test.py -- re-run those first if the classifier or dataset changed.
"""

import json
from pathlib import Path

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).parent
OUT = HERE.parent / "Group_Assignment_2.docx"

GREEN = RGBColor(0x1A, 0x7F, 0x37)
RED = RGBColor(0xB3, 0x26, 0x1E)
AMBER = RGBColor(0x8A, 0x6D, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)


# ── helpers ──────────────────────────────────────────────────────────────────
def shade(cell, hexfill):
    el = docx.oxml.parse_xml(
        r'<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(
            r'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
            hexfill))
    cell._tc.get_or_add_tcPr().append(el)


def runs(p, parts, size=10):
    """parts: str, or list of (text, **opts) tuples/dicts."""
    if isinstance(parts, str):
        parts = [(parts, {})]
    for text, opt in parts:
        r = p.add_run(text)
        r.font.size = Pt(opt.get("size", size))
        r.bold = opt.get("b", False)
        r.italic = opt.get("i", False)
        if opt.get("c"):
            r.font.color.rgb = opt["c"]
        if opt.get("mono"):
            r.font.name = "Consolas"
    return p


def para(doc, parts, size=10, after=5, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align:
        p.alignment = align
    return runs(p, parts, size)


def head(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(13 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13 if level == 1 else 11)
    r.font.color.rgb = BLACK
    return p


def bullet(doc, parts, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    return runs(p, parts, size)


def table(doc, headers, rows, widths, center_from=99, highlight=(), size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False

    def fill(cell, content, w, bold=False, align=None, fillc=None):
        cell.width = Inches(w)
        if fillc:
            shade(cell, fillc)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        if align:
            p.alignment = align
        parts = [(content, {})] if isinstance(content, str) else content
        for text, opt in parts:
            r = p.add_run(text)
            r.font.size = Pt(opt.get("size", size))
            r.bold = opt.get("b", bold)
            r.italic = opt.get("i", False)
            if opt.get("c"):
                r.font.color.rgb = opt["c"]

    for i, h in enumerate(headers):
        fill(t.rows[0].cells[i], h, widths[i], bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, fillc="E8E8E8")
    t.rows[0]._tr.get_or_add_trPr().append(
        docx.oxml.parse_xml(r'<w:tblHeader {}/>'.format(
            r'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"')))

    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, c in enumerate(row):
            fill(cells[i], c, widths[i],
                 align=WD_ALIGN_PARAGRAPH.CENTER if i >= center_from else None,
                 fillc="FBEAEA" if ri in highlight else None)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


B = {"b": True}
I = {"i": True}


# ═════════════════════════════════════════════════════════════════════════════
doc = docx.Document()
s = doc.sections[0]
s.page_width, s.page_height = Inches(8.5), Inches(11)
s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0.75)
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

para(doc, [("Group Assignment 2 — Evaluation & Measurement", {"b": True, "size": 15})], after=2)
para(doc, [("Energy Anomaly Explainer · AI Analytics Aces · BUSN 43800 · "
            "University of Chicago Booth", {"c": GREY, "size": 8.5})], after=10)

# ───────────────────────────── PART A ────────────────────────────────────────
head(doc, "5.1  Part A — The core task specification")
para(doc, [
    ("The task. ", B),
    ("Given one spike the detector has already flagged, name its most likely cause and tell the "
     "operations manager what to do about it. One spike in, one record out. Detection, the "
     "statistical significance test and the dispatch gate are separate components and are not "
     "part of this task.", {}),
])

head(doc, "2.1  Input", 2)
for x in [
    "Accepts hourly kWh electricity consumption data (smart meter or BMS export), via CSV upload in Settings",
    "Outdoor temperature from a public weather API (Open-Meteo), fetched automatically — no sensor required",
    "No additional sensors, sub-metering, or new infrastructure",
    "Every import is checked against the existing database before use — unit scale, negatives, runs "
    "of zeros, coverage gaps, unregistered sub-systems, and temperatures that read as Celsius in a "
    "Fahrenheit column. Findings are prompts to check, never rejections.",
]:
    bullet(doc, x)

head(doc, "2.2  Classification engine", 2)
para(doc, "Classifies each flagged spike into one of three top-level categories, and one of 14 "
          "subtypes within them:")
for x in [
    "Equipment fault — 7 subtypes (compressor fault, refrigerant leak, HVAC fan failure, HVAC "
    "filter blockage, lighting control fault, door seal failure, power surge)",
    "Operational variation — 4 subtypes (peak throughput day, unscheduled overtime, weather-driven "
    "HVAC surge, temporary equipment rental)",
    "Data anomaly — 3 subtypes (meter dropout, sensor noise spike, communication error)",
]:
    bullet(doc, x)
para(doc, "Evidence given to the model:", before=4)
for x in [
    "The spike's peak, its baseline for that hour and temperature, the excess in kWh and as a multiple",
    "Duration, outdoor temperature at detection, timestamp with weekday",
    "The sub-system's name and type",
    "The consumption curve for three hours either side of the spike",
    "The complete 14-item catalogue, identical on every case",
]:
    bullet(doc, x)
para(doc, [
    ("Deliberately withheld: ", B),
    ("the ground-truth label, any other anomaly, any previous classification, and the outcome of "
     "the statistical test. Held-out labels are opened only by the scoring script, after every "
     "prediction is written to disk — which is why the measured precision is not circular.", {}),
], before=4)

head(doc, "2.3  Output", 2)
para(doc, "Seven fields, enforced by the API so a malformed record cannot be returned:")
for x in [
    "Top-level class and subtype",
    "Confidence, as a genuine posterior rather than a default",
    "Plain-English explanation, written for a non-technical reader",
    "A specific recommended next action — dispatch a technician, monitor for recurrence over the "
    "next 24 hours, or dismiss as a known operational event",
    "The manager's next step, grounded in this spike: which trade to send, which hour to watch, or "
    "which artefact to log",
    "The symptom for the technician, in an engineer's words — required when dispatching, empty otherwise",
]:
    bullet(doc, x)
para(doc, [
    ("When none of the three actions fit. ", B),
    ("Three actions cannot cover every spike, and a manager forced to pick the closest wrong one "
     "leaves no trace that the taxonomy failed. Every anomaly therefore carries a fifth path: "
     "record an ", {}),
    ("exception", I),
    (" with a free-text reason. Those notes are the shortlist for what the catalogue is missing — "
     "both misclassifications in Part C would have surfaced through it long before anyone scored "
     "a rubric.", {}),
], before=4)

head(doc, "2.4  Response time", 2)
bullet(doc, "Bar: classification and recommended action within 5 minutes of upload or query")
bullet(doc, [("Measured: 10.0 s average, 17.8 s worst case", B)])

head(doc, "3.  Performance specifications")
head(doc, "3.1  Primary bar — classification accuracy", 2)
table(doc, ["Metric", "Bar", "Result", "Verdict"], [
    ["Precision, equipment fault class", "≥ 75%", "82%  (9 TP, 2 FP)", [("PASS", {"b": True, "c": GREEN})]],
    ["Recall, equipment fault class", "≥ 70%", "100%  (0 FN)", [("PASS", {"b": True, "c": GREEN})]],
], [2.5, 0.8, 1.8, 0.9], center_from=1)
para(doc, [("Validated against a 14-case held-out synthetic test set with known ground-truth "
            "labels. We planned 15 and report 14 — one injected fault was never detected, and a "
            "case the detector does not surface cannot be classified.", {"c": GREY, "size": 8.5})])

head(doc, "3.2  Secondary bar — explanation quality", 2)
for x in [
    "Four yes/no points per output — names a cause, states the evidence, gives a specific next "
    "step, avoids jargon. All four required for a case to pass",
    "Two human raters score independently, blind to the label and to each other; raw agreement "
    "reported, splits go to a third reader",
]:
    bullet(doc, x)
bullet(doc, [("Status: an LLM pre-rating scores 13/14. Human scoring outstanding — the judge is "
              "not authoritative until it has been shown to agree with the readers it replaces.", B)])

head(doc, "3.3  Failure condition", 2)
para(doc, "If precision on the equipment fault class falls below 60%, the product must be "
          "redesigned to require shift schedule data as a mandatory input. If fault vs. "
          "operational variation confusion persists while precision remains above 60%, the "
          "redesign is structural rather than input-side: detection and classification move from "
          "the sub-system to the facility, with the sub-system carried as an attribute, and the "
          "alert threshold is restated relative to each meter's own variability rather than in "
          "absolute kWh. Adding inputs is not a remedy for a confusion produced by the unit of "
          "analysis.")
para(doc, [
    ("Amended after testing. ", {"b": True, "size": 8.5, "c": GREY}),
    ("The original clause prescribed shift schedule ", {"size": 8.5, "c": GREY}),
    ("or", {"i": True, "size": 8.5, "c": GREY}),
    (" outdoor temperature. Temperature was already a mandatory input before the trigger fired, "
     "and shift schedule was implemented and measured as blocked — see Part C. The second limb of "
     "the original condition has been triggered; this states what the remedy should be, not "
     "whether the condition was met.", {"size": 8.5, "c": GREY}),
])

head(doc, "4.  User & deployment specifications")
head(doc, "4.1  Target user", 2)
bullet(doc, "Operations managers at U.S. distribution centers ≥100,000 sq ft")
bullet(doc, "Assumed to lack the statistical background to interpret raw time-series data")
head(doc, "4.2  Deployment context", 2)
bullet(doc, "Facilities must have access to hourly smart meter or BMS consumption data")
bullet(doc, "Facilities must experience at least 2 unexplained spikes per month in the historical record")
head(doc, "4.3  Constraints", 2)
for x in [
    "Must operate on hourly kWh readings already available to the manager, plus outdoor "
    "temperature from a public weather API",
    "No new sensors, sub-metering, or infrastructure spend. Any further input must be obtainable "
    "without installation or procurement",
    "Must not require engineering consultation for the initial triage decision",
]:
    bullet(doc, x)
para(doc, [
    ("Amended. ", {"b": True, "size": 8.5, "c": GREY}),
    ("The original clause read “solely on hourly kWh readings”. That was never true of the build: "
     "temperature is carried on every reading, drives the agent's baseline, and appears in every "
     "prompt. Removing it disables the weather-adjusted baseline and the entire "
     "not-statistically-significant path. The constraint that actually binds is no capital "
     "expenditure and nothing to install.", {"size": 8.5, "c": GREY}),
])

head(doc, "5.  Data specifications")
head(doc, "5.1  Training / validation dataset (synthetic)", 2)
bullet(doc, "12-month hourly dataset — 78,840 readings across 3 facilities and 9 sub-systems")
bullet(doc, [("Outdoor temperature is real", B),
             (" — hourly observations for Chicago, Milwaukee and Indianapolis from the "
              "Open-Meteo ERA5 archive. Consumption is computed from it, so weather adjustment "
              "is tested rather than assumed", {})])
for x in [
    "Two shifts: day 06:00–14:00, night 22:00–06:00; mixed dry and refrigerated storage",
    "Base load 40–60 kW; forklift charging bursts post-shift; 4-hour refrigeration compressor "
    "cycles; weather-correlated HVAC load; stochastic variation throughout",
    "Reproducible from generate_dataset.py — the original workbook shipped without the script "
    "that made it",
]:
    bullet(doc, x)
head(doc, "5.2  Injected anomaly classes (with known ground-truth labels)", 2)
for x in [
    "Anomalies are injected, then found by a real detector. detected_at, spike_kwh and "
    "baseline_kwh are measured off the series, not asserted",
    "An injection the detector misses is reported as a detection miss rather than silently "
    "becoming a test case",
    "Three strata: core (14 single-cause cases, the success bar), co-occurring (2 causes in one "
    "hour), sub-threshold (deliberately marginal spikes)",
    "One ground-truth source. The previous workbook carried two that disagreed, and its own "
    "documentation prescribed the one covering 4 of 15 cases",
]:
    bullet(doc, x)

head(doc, "6.  Integration specifications")
bullet(doc, [("AI explanation layer powered by ", {}),
             ("Claude Opus 5 (claude-opus-5)", B),
             (" via the Anthropic Messages API, with the output constrained by a JSON schema so "
              "a malformed record cannot be returned", {})])
bullet(doc, "Streamlit web application for live presentation and user interaction, four screens")
bullet(doc, "A four-tool agent orchestrator runs when the Classification Panel opens: fetch "
            "readings, compute baseline, run significance test, fetch comparable events")

head(doc, "7.  Scope limitations")
bullet(doc, "Results validated at the classification level on synthetic data only")
bullet(doc, "Real-world validation requires field deployment with maintenance log access, out of "
            "scope for this phase")
bullet(doc, [("The detector raises 562 alarms across the year; only the 25 matching an injected "
              "anomaly reach the test set. Every case scored is therefore a genuine planted "
              "event, so precision measures label accuracy and never the harder question of "
              "telling a real event from a false alarm", {"c": RED})])

head(doc, "A good output and a bad one", 2)
table(doc, ["", "Good — ANO-2010", "Bad — ANO-2009"], [
    ["What it did",
     "Called a midnight HVAC spike Weather-Driven HVAC Surge at 0.72, action dismiss. Correct.",
     "Called an eight-hour refrigeration spike Compressor Fault at 0.78, action dispatch. It was "
     "a Peak Throughput Day."],
    ["Why",
     "Reasoned from the shape — the flagged hour sits below the hours before it and the evening "
     "tails off as the air cools — and named the mechanism rather than the magnitude.",
     "The reasoning is sound and the evidence is absent. A sustained overnight load that does not "
     "ease as it cools does look like a compressor fault when shift schedules are not an input."],
    ["Next step",
     "“Log this as expected warm-night cooling load on HVAC Unit 1… so the same midnight hour is "
     "not re-flagged next month.” Specific, and closes the loop.",
     "Sends a technician overnight to a healthy compressor bank. Confidently wrong is worse than "
     "uncertain: at 0.78 it clears the commit gate."],
], [0.75, 2.65, 2.6])
para(doc, [("Full prompt text in Appendix 1.", {"c": GREY, "size": 8.5})])

# ───────────────────────────── PART B ────────────────────────────────────────
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
head(doc, "5.2  Part B — The evaluation rubric")
para(doc, [("One output = one classified spike. ", B),
           ("Each of the five criteria is scored on every case; a record passes only if it passes "
            "all five. Aggregates are rolled up from those per-case scores, never measured "
            "directly.", {})])

head(doc, "The per-output rubric", 2)
table(doc, ["Criterion", "How it is judged", "Acceptable at", "Result"], [
    ["Correctness",
     "Does the record name the right cause? Top-level class and subtype each compared to the "
     "held-out label by exact match.", "Both exact",
     [("86% class", {"b": True, "c": RED}), ("  ", {}), ("50% subtype", {"b": True, "c": RED})]],
    ["Completeness",
     "Are all seven fields present and non-empty, with a symptom supplied whenever the action is "
     "dispatch?", "All fields", [("14/14", {"b": True, "c": GREEN})]],
    ["Calibration",
     "Does confidence agree with correctness — commit (≥0.75) when right, hedge (<0.75) when "
     "wrong? A wrong label at ≥0.75 is an overclaim and fails the case outright.",
     "No overclaims; ≥80% agree",
     [("43%", {"b": True, "c": RED}), ("  ", {}), ("1 overclaim", {"b": True, "c": RED})]],
    ["Usability",
     "Could the manager act unaided? Four yes/no points — names a cause, states the evidence, "
     "gives a specific next step, avoids jargon. Two human raters, blind.", "≥80% of cases",
     [("13/14 judge", {"b": True, "c": AMBER}), ("  ", {}), ("humans pending", {"b": True, "c": AMBER})]],
    ["Timeliness", "Wall clock from request to returned record, timed in the classifier.", "< 5 min",
     [("10.0 s avg", {"b": True, "c": GREEN}), ("  ", {}), ("17.8 s max", {"b": True, "c": GREEN})]],
], [0.95, 3.0, 1.1, 1.1], center_from=2, highlight=(0, 2, 3))
para(doc, [("Scored on the 14-case core stratum (9 equipment faults), classified by "
            "claude-opus-5. Per-case scores in Appendix 2. Reproduce with python3 "
            "score_rubric.py.", {"c": GREY, "size": 8.5})])

head(doc, "Aggregates rolled up from the per-case scores", 2)
table(doc, ["Aggregate", "How it is derived", "Acceptable at", "Result"], [
    ["Precision, equipment fault",
     "Share of predicted faults that are faults — from Correctness. 9 TP, 2 FP.", "≥ 75%",
     [("82%", {"b": True, "c": GREEN})]],
    ["Recall, equipment fault",
     "Share of true faults caught; an abstention counts as a miss. 0 FN.", "≥ 70%",
     [("100%", {"b": True, "c": GREEN})]],
    ["Decision value",
     "Cost of following the tool vs. two fixed policies at $300 a dispatch and $2,000 a missed "
     "fault, over all 25 anomalies.", "beat $7,500", [("$35,200", {"b": True, "c": RED})]],
], [1.35, 2.75, 1.0, 1.05], center_from=2, highlight=(2,))

head(doc, "Do we use a model to judge outputs?", 2)
para(doc, [("Four of the five criteria use no judge at all. ", B),
           ("Correctness is an exact string match against held-out labels; Completeness and "
            "Calibration are field and threshold checks; Timeliness reads a timer. Nothing in "
            "that path can drift, and the model is never both author and marker.", {})])
para(doc, [("Usability is the exception, and we checked the judge rather than trusting it. ", B),
           ("Two human raters score independently, blind to the label and to each other. The "
            "rating pack carries only the explanation, the action and the symptom — no class, no "
            "confidence, no ground truth — so a rater cannot be swayed by knowing which cases "
            "were right, and the instructions say plainly that they are scoring how it reads, "
            "not whether it is correct. Raw agreement is reported per point and splits go to a "
            "third reader.", {})])
para(doc, [("An LLM judge scored the same 14 cases at 13/14, failing ANO-2008 on ", {}),
           ("names a cause", I),
           (" — “a mechanical problem in the air-handling side” is a category, not a mechanism. "
            "We do not quote that as the result. The scoring script prints it as not "
            "authoritative and refuses to stand on it until human scores exist, because a judge "
            "is usable only once it has been shown to agree with the readers it replaces.", {})])

# ───────────────────────────── PART C ────────────────────────────────────────
head(doc, "5.3  Part C — The test set")
para(doc, [("14 cases. ", B),
           ("The core stratum of a 12-month generated dataset: 9 equipment faults, 3 operational "
            "variations, 2 data anomalies, across three facilities and all nine sub-systems. Two "
            "further strata — 5 co-occurring causes and 6 deliberately marginal spikes — are held "
            "out of the headline number and reported separately.", {})])

head(doc, "Results", 2)
table(doc, ["Criterion", "Result", "Verdict"], [
    ["Correctness — class", "12 / 14 (86%)", [("FAIL", {"b": True, "c": RED})]],
    ["Correctness — subtype", "7 / 14 (50%)", [("FAIL", {"b": True, "c": RED})]],
    ["Completeness", "14 / 14", [("PASS", {"b": True, "c": GREEN})]],
    ["Calibration", "43%, 1 overclaim", [("FAIL", {"b": True, "c": RED})]],
    ["Timeliness", "10.0 s avg", [("PASS", {"b": True, "c": GREEN})]],
    ["Precision / Recall", "82% / 100%", [("PASS", {"b": True, "c": GREEN})]],
    ["Decision value", "$35,200 vs $7,500", [("FAIL", {"b": True, "c": RED})]],
], [2.1, 2.1, 1.1], center_from=1, highlight=(0, 1, 3, 6))

head(doc, "Every failure, and what it was", 2)
table(doc, ["Failure", "What went wrong", "Diagnosis"], [
    ["ANO-2009, ANO-2011",
     "Peak Throughput Day and Temporary Equipment Rental both read as equipment faults, at 0.78 "
     "and 0.72. Neither shift schedules nor rental logs are inputs the product receives, so no "
     "prompt can separate a busy shift from a stuck compressor on kWh alone.",
     [("Specification", B)]],
    ["5 of 14 cases",
     "Right top-level class, wrong subtype — e.g. Door Seal Failure read as Refrigerant Leak. "
     "Hourly kWh underdetermines which fault; two mechanisms with the same load signature are "
     "indistinguishable at this sampling rate.", [("Evidence", B)]],
    ["1 overclaim, 7 hedges",
     "Confidence averages 0.70 when right and 0.75 when wrong. The model's confidence is "
     "miscalibrated, and inverted rather than merely noisy. This is a property of the model, not "
     "of our inputs.", [("Model", B)]],
    ["Decision value",
     "82% precision and 100% recall, and following the tool still costs 4.7× dispatching on "
     "everything. We asked “what is this spike?” when the decision needs “should I dispatch?”. A "
     "correct label at 0.72 still produces no dispatch.", [("Question", B)]],
], [1.05, 3.55, 0.9], center_from=2)

head(doc, "Are the specification failures reducible? We tested it", 2)
para(doc, "Our failure condition names this confusion as the trigger to require shift schedule "
          "data. We implemented that remedy and a second one, and measured both.")
table(doc, ["Remedy", "Result", "Why"], [
    [[("Shift schedule", B), ("  ", {}), ("the prescribed remedy", I)],
     [("Blocked", {"b": True, "c": RED}), ("  ", {}), ("precision 82% → 64%", {})],
     "A day planned at 130% of normal adds 7–10 kWh to a sub-system; the alarm fires at 20. A "
     "realistic busy day never reaches the detector, so any operational event large enough to "
     "alarm is larger than its own explanation — at which point “fault” is the correct read."],
    [[("Sibling co-movement", B), ("  ", {}), ("needs no new input", I)],
     [("Net negative", {"b": True, "c": AMBER}), ("  ", {}), ("recall 100% → 78%", {})],
     "A fault lifts one meter and site activity lifts every meter, so this fixed the Peak "
     "Throughput case and un-inverted calibration. But it cannot tell one shared cause from two "
     "simultaneous single-asset events, and a real refrigerant leak was read as operational."],
], [1.15, 1.15, 3.2], center_from=1)
para(doc, "Neither dead end is about which variables the model receives. Both are about the unit "
          "of analysis — nine meters treated as nine independent problems, when operational "
          "variation is a facility-level phenomenon — and the alert threshold, set per sub-system "
          "in absolute kWh. Under the specification's own 167:1 miss-to-false-alarm ratio, "
          "trading recall for precision is the wrong direction, so neither was merged. Both are "
          "preserved on branches with their measurements.", before=4)

head(doc, "Where the $35,200 comes from", 2)
table(doc, ["Policy", "Sent", "Caught", "Dispatch spend", "Missed-fault exposure", "Total"], [
    ["Follow the tool", "4", "3 / 20", "$1,200", "17 × $2,000 = $34,000",
     [("$35,200", {"b": True, "c": RED})]],
    ["Dispatch on every spike", "25", "20 / 20", "$7,500", "none",
     [("$7,500", {"b": True, "c": GREEN})]],
    ["Dispatch on none", "0", "0 / 20", "$0", "20 × $2,000 = $40,000",
     [("$40,000", {"b": True, "c": RED})]],
], [1.4, 0.5, 0.7, 1.0, 1.6, 0.9], center_from=1, highlight=(0,))
para(doc, [("97% of the tool's cost is exposure, not spend. ", B),
           ("It disburses $1,200 and leaves $34,000 of faults uninvestigated. Two multipliers "
            "drive it: a miss costs 6.7× a dispatch, and 20 of these 25 anomalies are genuine "
            "faults. At an 80% base rate you would need to be almost certain before ", {}),
           ("not", I),
           (" sending someone, and a gate set at 0.75 confidence is nowhere near that.", {})],
     before=4)
para(doc, [("One modelling choice to state plainly: a ", {"size": 9, "c": GREY}),
           ("monitor", {"i": True, "size": 9, "c": GREY}),
           (" counts as a miss, because nobody is dispatched. If monitoring reliably catches the "
            "fault later at reduced cost, the figure falls. We have no evidence either way, so we "
            "scored the conservative reading.", {"size": 9, "c": GREY})])
para(doc, [("The calibration failure is the one we did not anticipate, and it matters most. ", B),
           ("The 0.75 commit gate selects for errors: it admits the single overclaim while "
            "hedging seven correct calls. Only 4 of 25 spikes clear it, so 17 of 20 real faults "
            "are told to “monitor” and nobody is dispatched. A rubric that stopped at precision "
            "and recall would have reported success.", {})])

# ───────────────────────────── PART D ────────────────────────────────────────
head(doc, "5.4  Part D — The measurement layer")
para(doc, "Our inputs are numeric, not unstructured text, so this part covers both what the "
          "assignment asks of a measurement layer and where our inputs come from.")

head(doc, "The construct", 2)
para(doc, [("The cause of a flagged spike", B),
           (" — a latent categorical variable taking one of 14 values in three families, together "
            "with ", {}),
           ("confidence", B),
           (", a claimed posterior probability that the assigned cause is correct. The conversion "
            "is from a 7-hour window of hourly kWh plus temperature and timing into that pair. "
            "Confidence is the more demanding claim: a category can be checked against a label, "
            "whereas a probability claim is only meaningful if it is calibrated — which is why "
            "Calibration is scored as its own criterion rather than folded into Correctness.", {})])

head(doc, "Worked examples", 2)
table(doc, ["Input signal", "→ Construct", "Correct?"], [
    ["HVAC Unit 1, midnight, 21.4 kWh above a 38.4 baseline, 76°F, one hour; the flagged hour "
     "sits below the three hours before it and the evening tails off.",
     "operational_variation / Weather-Driven HVAC Surge, 0.72, dismiss",
     [("Yes", {"b": True, "c": GREEN})]],
    ["Compressor Bank, 20:00, roughly double normal draw held for eight hours without settling, "
     "outside air cooling.", "equipment_fault / Compressor Fault, 0.78, dispatch",
     [("No — was a Peak Throughput Day", {"b": True, "c": RED})]],
    ["HVAC Unit 2, 14:00, 116.8 kWh above baseline for a single hour, ~6× the unit's normal draw, "
     "neighbours unremarkable.", "data_anomaly / Sensor Noise Spike, 0.78, dismiss",
     [("Yes", {"b": True, "c": GREEN})]],
], [2.9, 1.9, 1.2], center_from=2)

head(doc, "Error — where the measure is wrong, and how we know", 2)
para(doc, "Against held-out labels the top-level class is right on 12 of 14 and the subtype on 7 "
          "of 14. So the family of cause is measured reasonably well and the specific mechanism "
          "is close to a coin flip. Confidence is worse than uninformative: it averages 0.70 when "
          "the label is right and 0.75 when it is wrong, so it runs against correctness. We know "
          "this because the labels were generated by a documented injection process and opened "
          "only after every prediction was written.")

head(doc, "Stability — same input, three runs", 2)
para(doc, [("The top-level class, the subtype and the recommended action were ", {}),
           ("identical on all 14 cases across all three runs", B),
           (". Only confidence moved, by 0.041 on average and 0.10 at worst. The measure is "
            "reproducible; its errors are systematic rather than random, which means they will "
            "not average out over more cases and cannot be fixed by sampling more.", {})])

head(doc, "Bias", 2)
table(doc, ["Group", "Accuracy", "Group", "Accuracy"], [
    ["Chicago South DC", "4/5 = 80%", "HVAC", "5/6 = 83%"],
    ["Indianapolis East", "3/4 = 75%", "Lighting", "3/4 = 75%"],
    ["Milwaukee Central", "4/5 = 80%", "Refrigeration", "3/4 = 75%"],
    ["Warm months", "6/8 = 75%", "Cold months", "5/6 = 83%"],
], [1.55, 1.0, 1.55, 1.0], center_from=1)
para(doc, [("Facility, system type and season are flat within the noise of a 14-case set. ", {}),
           ("The bias is in the true class:", B),
           (" equipment faults 8/9 (89%), data anomalies 2/2 (100%), operational variation ", {}),
           ("1/3 (33%)", B),
           (". The measure resolves ambiguity toward “fault”. Given a miss costs far more than a "
            "needless visit that is the safer direction to err, but it is not free — it is the "
            "same asymmetry that produces the two specification failures in Part C, and it means "
            "any deployment will over-report faults during peak season.", {})], before=4)

head(doc, "Where the inputs come from", 2)
para(doc, [("Hourly consumption is generated by a documented model (base load, shift ramps, "
            "four-hour compressor cycles, forklift charging, weather response, stochastic "
            "noise). ", {}),
           ("Outdoor temperature is real", B),
           (" — hourly observations for Chicago, Milwaukee and Indianapolis from the Open-Meteo "
            "archive — and consumption is computed from it, so weather adjustment is tested "
            "rather than assumed. Anomalies are injected and then found by a detector, so the "
            "recorded spike magnitude and baseline are measured from the series rather than "
            "asserted.", {})])

head(doc, "Protecting the measure from a bad feed", 2)
para(doc, "Every figure above assumes the inputs are what they claim to be. A file in watts "
          "rather than kilowatts, a meter that stopped reporting, or a sub-system code that does "
          "not exist all render a perfectly plausible dashboard while every classification behind "
          "it is meaningless — and none of it shows up in precision or recall, because the labels "
          "still match. Every import is profiled against the settled history and checked for unit "
          "scale, negatives, runs of zeros, values beyond anything recorded, duplicate "
          "timestamps, coverage gaps, unregistered sub-systems, unreadable or future timestamps, "
          "and temperatures that read as Celsius in a Fahrenheit column. That last one needs a "
          "comparison against the range the database has actually seen — 70°F becomes 21°C, which "
          "passes any plausible-range check.")

head(doc, "Why this does not fully identify what we need", 2)
para(doc, "The construct is the cause of a spike, and cause is only recoverable from consumption "
          "when the operational context that would explain it is absent. It is not: shift "
          "schedules, throughput and equipment rentals all move load and none is an input. The "
          "two Part C specification failures are that gap, measured.")

# ───────────────────────────── APPENDICES ────────────────────────────────────
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
head(doc, "Appendix 1 — The prompt")
para(doc, [("System prompt, verbatim. Sent unchanged on every case.", {"c": GREY, "size": 8.5})])
for line in (HERE / "appendix_prompt.txt").read_text().split("\n"):
    if line.strip():
        p = para(doc, [(line.rstrip(), {"size": 7.5, "mono": True})], after=0)
        if line.startswith((" ", "-", "\t")):
            p.paragraph_format.left_indent = Inches(0.18)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

head(doc, "Output schema", 2)
para(doc, [("top_level_class (enum: equipment_fault | operational_variation | data_anomaly) · "
            "classification_type_id · confidence_score (0–1) · explanation_text · "
            "recommended_action (enum: dispatch | monitor | dismiss) · next_action · "
            "symptom_to_check. All seven required; no additional properties permitted.",
            {"size": 7.5, "mono": True})])

head(doc, "Appendix 2 — Per-case scores, core stratum")
cases = json.loads((HERE / "appendix_cases.json").read_text())
table(doc, ["Case", "True", "Predicted", "Class", "Subtype", "Conf.", "Action", "Sec"],
      [[c if i < 3 else [(c, B)] for i, c in enumerate(row)] for row in cases],
      [0.6, 0.85, 0.95, 0.6, 0.7, 0.6, 0.85, 0.5], center_from=3,
      highlight=tuple(i for i, r in enumerate(cases) if r[3] == "N"))
para(doc, [("Reproduce every figure in this document with python3 score_rubric.py, python3 "
            "stability_test.py --report-only, and python3 usability_test.py --score.",
            {"c": GREY, "size": 8.5})])

para(doc, [("Honor Code. We pledge our honor that we have not violated the Honor Code in "
            "preparation of this case assignment / group Project.", {"c": GREY, "size": 8.5})],
     before=8)
para(doc, [("AI-use disclosure. Claude (Anthropic) is the classifier under evaluation and was "
            "also used to implement the remedies tested in Part C, to compute the measurements, "
            "and to draft this document. All criteria, thresholds, diagnoses and interpretations "
            "are the team's own. Every figure is reproducible from the repository.",
            {"c": GREY, "size": 8.5})])

doc.save(OUT)
print("wrote", OUT)
