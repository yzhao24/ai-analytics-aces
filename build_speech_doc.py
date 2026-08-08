"""
The spoken script as cue cards, one page per speaker.

    python3 build_speech_doc.py

Reads speaker_notes.py, so it is always the same words that are in the deck's
notes pane. Print it, or read it off a phone.
"""

from pathlib import Path

import docx
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from speaker_notes import COACHING, SCRIPT, TIMING, WPM

OUT = Path(__file__).parent.parent / "Final_Presentation_Speech.docx"
GREY = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GOLD = RGBColor(0x8A, 0x6D, 0x00)

SLIDE_TITLES = {
    1: "Title", 2: "The problem, as a measurable gap",
    3: "We committed to a test before we built anything", 4: "What we built",
    5: "LIVE DEMO", 6: "How we decide an output is good",
    7: "What the test set says — including the failures",
    8: "The failure that matters", 9: "What it cannot do",
    10: "What happens next", 11: "Where we landed",
}

# what to say if the question comes, keyed to the slide it follows from
QA = [
    ("“Isn't 82% precision good enough to ship?”",
     "No — and that is the point of slide 8. Precision is a property of the label. "
     "Dispatch is a property of the gate on top of it. Ours sends four technicians "
     "out of twenty-five spikes, so 82% precision never converts into a decision. "
     "Fix the gate and the same classifier becomes useful."),
    ("“Why is your test set only 14 cases?”",
     "That is the core stratum. There are 25 anomalies in total; 11 are held out as "
     "co-occurring or deliberately marginal and reported separately. We planned 15 "
     "core cases and report 14 because one injected fault was never detected — and "
     "a case the detector does not surface cannot be classified."),
    ("“How do you know the model isn't grading itself?”",
     "Four of the five criteria involve no model. The ground-truth labels sit in a "
     "sheet the classifier never receives, and the scoring script opens them only "
     "after every prediction is written to disk."),
    ("“Would real data change your results?”",
     "It would change the error rate, and we cannot say in which direction. What it "
     "would not change is the specification failure: shift schedules are not an "
     "input in the real world either. That gap is structural, not synthetic."),
    ("“Why not just use a bigger threshold / better prompt?”",
     "We tried both remedies our own spec prescribed and measured them failing — "
     "shift schedule dropped precision to 64%, sibling co-movement dropped recall "
     "to 78%. The constraint is the unit of analysis, not the prompt."),
    ("“What does the exception button actually do?”",
     "It records a free-text reason and tags the anomaly as an exception. It is our "
     "cheapest source of information about what the 14-cause taxonomy is missing."),
]


def main():
    doc = docx.Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.85)
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)
    n.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    def para(parts, size=11, after=6, before=0, color=BLACK, align=None, line=1.25):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.line_spacing = line
        if align:
            p.alignment = align
        for text, o in ([(parts, {})] if isinstance(parts, str) else parts):
            r = p.add_run(text)
            r.font.size = Pt(o.get("size", size))
            r.bold = o.get("b", False)
            r.italic = o.get("i", False)
            r.font.color.rgb = o.get("c", color)
        return p

    total = sum(sec for _, sec in TIMING.values())
    para([("Final Presentation — Speaking Script", {"b": True, "size": 18})], after=2)
    para([(f"Energy Anomaly Explainer · AI Analytics Aces · "
           f"{total // 60}:{total % 60:02d} across 11 slides · {WPM} words/min",
           {"c": GREY, "size": 9.5})], after=12)

    # running order
    para([("Running order", {"b": True, "size": 13})], after=5)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for i, (h, w) in enumerate([("Slide", 0.6), ("Speaker", 1.2), ("What it covers", 3.9),
                                ("Time", 0.7)]):
        c = t.rows[0].cells[i]
        c.width = Inches(w)
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
    for i in sorted(TIMING):
        who, sec = TIMING[i]
        cells = t.add_row().cells
        for j, (v, w) in enumerate([(str(i), 0.6), (who, 1.2), (SLIDE_TITLES[i], 3.9),
                                    (f"{sec // 60}:{sec % 60:02d}", 0.7)]):
            cells[j].width = Inches(w)
            p = cells[j].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(v)
            r.font.size = Pt(10)
            r.bold = j == 1
    para([("Every member speaks. Slide 5 is performed, not read — the two minutes are "
           "the demo itself.", {"c": GREY, "size": 9.5})], before=6, after=0)

    # one page per speaker
    order, seen = [], set()
    for i in sorted(TIMING):
        who = TIMING[i][0]
        if who not in seen:
            seen.add(who)
            order.append(who)
    for who in order:
        slides = [i for i in sorted(TIMING) if TIMING[i][0] == who]
        secs = sum(TIMING[i][1] for i in slides)
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        para([(who, {"b": True, "size": 20})], after=1)
        para([(f"Slides {', '.join(str(i) for i in slides)}  ·  "
               f"{secs // 60}:{secs % 60:02d} total", {"c": GREY, "size": 10})], after=10)
        for i in slides:
            sec = TIMING[i][1]
            para([(f"SLIDE {i}   ", {"b": True, "size": 10, "c": GOLD}),
                  (SLIDE_TITLES[i], {"b": True, "size": 12}),
                  (f"   ({sec // 60}:{sec % 60:02d})", {"c": GREY, "size": 10})],
                 before=8, after=4)
            if COACHING.get(i):
                for block in COACHING[i].split("\n\n"):
                    if block.strip():
                        para([(" ".join(block.split()) if not block.startswith("      ")
                               else block.strip(), {"i": True, "size": 10.5, "c": GREY})],
                             after=5, line=1.3)
            for block in SCRIPT[i].split("\n\n"):
                if block.strip():
                    para(" ".join(block.split()), size=12, after=7, line=1.4)

    # Q&A prep
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    para([("Questions we should expect", {"b": True, "size": 20})], after=2)
    para([("Whoever owns the slide the question lands on answers it. Short answers — "
           "the honest ones are already in the deck.", {"c": GREY, "size": 10})], after=10)
    for q, a in QA:
        para([(q, {"b": True, "size": 12})], before=7, after=3)
        para(a, size=11.5, after=4, line=1.35)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    para([("Before you walk in", {"b": True, "size": 20})], after=8)
    for item in [
        "Run the demo end-to-end on the presentation machine — the assignment checklist "
        "requires it, and first load takes several seconds while 78,840 rows cache.",
        "python3 run_dashboard.py, then leave it open on the Dashboard screen.",
        "Rebuild sample_bad_export.csv (the one-liner is in the README) and leave it on "
        "the desktop.",
        "Have the fallback screen recording open in another window. If you use it, say "
        "out loud that it is a recording.",
        "Submit on Canvas: the repository link, the one-page product brief, and these "
        "slides. The brief carries the Honor Code pledge and the AI-use disclosure.",
        "Decide who fields questions — default is whoever owns the slide it lands on.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.3
        r = p.add_run(item)
        r.font.size = Pt(11.5)

    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    main()
